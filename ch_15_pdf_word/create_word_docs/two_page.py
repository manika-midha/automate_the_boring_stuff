import docx

doc = docx.Document()

doc.add_paragraph('This is on the first page !')
doc.add_page_break()
doc.add_paragraph('This is on the second page !')

doc.save('twopage.docx')