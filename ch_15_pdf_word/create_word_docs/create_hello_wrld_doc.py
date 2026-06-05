import docx

doc = docx.Document()
doc.add_paragraph('1st para : Hello World!', 'Title')
para_obj1 = doc.add_paragraph('2nd para : How are you doing !')
para_obj2 = doc.add_paragraph('3rd para : How is the weather ?')

para_obj1.add_run('This text is being added to the 2nd para.')

doc.save('helloworld_style.docx')
