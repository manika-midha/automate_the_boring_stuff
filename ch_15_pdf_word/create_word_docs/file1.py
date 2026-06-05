import docx

doc = docx.Document('demo.docx')
print(len(doc.paragraphs))

#print(doc.paragraphs)

# for obj in doc.paragraphs:
#     print(obj.text) #without styling
#
#     for run_obj in obj.runs :
#         print(run_obj.text)

print(doc.paragraphs[1].runs[0].text)