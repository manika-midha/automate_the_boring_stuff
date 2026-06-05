import docx

doc = docx.Document('demo.docx')

# 0 field
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'

# 1 field
print(doc.paragraphs[2].text)
print(doc.paragraphs[2].runs[0].text, doc.paragraphs[2].runs[1].text, doc.paragraphs[2].runs[2].text,
      doc.paragraphs[2].runs[3].text, sep='\n')

# print(doc.paragraphs[2].runs[0].text)
# print(doc.paragraphs[2].runs[1].text)
# print(doc.paragraphs[2].runs[2].text)
# print(doc.paragraphs[2].runs[3].text)

doc.paragraphs[2].runs[0].style = 'QuoteChar'
doc.paragraphs[2].runs[1].underline = True
doc.paragraphs[2].runs[3].underline = True

doc.save('restyled.docx')


